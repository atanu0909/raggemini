import os
import io
import tempfile
from typing import Dict, List, Optional
import streamlit as st

# Try multiple PDF libraries for better compatibility
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class DocumentProcessor:
    """Handles PDF and document processing for chapter extraction"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            st.error(f"Error reading PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error reading TXT: {str(e)}")
            return ""
    
    def extract_text_from_pdf_bytes(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF bytes with multiple library fallbacks"""
        
        # Method 1: Try PyPDF2 first
        if PYPDF2_AVAILABLE:
            try:
                text = self._extract_with_pypdf2(pdf_bytes)
                if text and text.strip():
                    return text
            except Exception as e:
                st.warning(f"PyPDF2 failed: {str(e)}")
        
        # Method 2: Try pypdf as fallback
        if PYPDF_AVAILABLE:
            try:
                text = self._extract_with_pypdf(pdf_bytes)
                if text and text.strip():
                    return text
            except Exception as e:
                st.warning(f"pypdf failed: {str(e)}")
        
        # Method 3: Try with temporary file as last resort
        try:
            text = self._extract_with_tempfile(pdf_bytes)
            if text and text.strip():
                return text
        except Exception as e:
            st.warning(f"Temporary file method failed: {str(e)}")
        
        # If all methods fail
        st.error("Could not extract text from PDF. The file might be corrupted or image-based.")
        return ""
    
    def _extract_with_pypdf2(self, pdf_bytes: bytes) -> str:
        """Extract text using PyPDF2"""
        pdf_stream = io.BytesIO(pdf_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                st.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                continue
        
        return text
    
    def _extract_with_pypdf(self, pdf_bytes: bytes) -> str:
        """Extract text using pypdf"""
        pdf_stream = io.BytesIO(pdf_bytes)
        pdf_reader = pypdf.PdfReader(pdf_stream)
        
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                st.warning(f"Error extracting page {page_num + 1}: {str(e)}")
                continue
        
        return text
    
    def _extract_with_tempfile(self, pdf_bytes: bytes) -> str:
        """Extract text using temporary file approach"""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(pdf_bytes)
            tmp_path = tmp_file.name
        
        try:
            if PYPDF2_AVAILABLE:
                with open(tmp_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            else:
                raise Exception("No PDF library available")
        finally:
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    def extract_text_from_docx_bytes(self, docx_bytes: bytes) -> str:
        """Extract text from DOCX bytes with robust error handling"""
        if not DOCX_AVAILABLE:
            st.error("python-docx library is not available")
            return ""
        
        try:
            docx_stream = io.BytesIO(docx_bytes)
            doc = docx.Document(docx_stream)
            
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                st.warning("No text could be extracted from the DOCX file.")
                return ""
            
            return text
            
        except Exception as e:
            st.error(f"Error reading DOCX: {str(e)}")
            return ""
    
    def process_uploaded_file(self, uploaded_file) -> Optional[str]:
        """Process uploaded file and extract text"""
        if uploaded_file is None:
            return None
        
        # Process file directly from memory without saving to disk
        try:
            # Extract text based on file extension
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()
            
            if file_extension == '.pdf':
                # Process PDF directly from bytes
                text = self.extract_text_from_pdf_bytes(uploaded_file.getvalue())
            elif file_extension == '.docx':
                # Process DOCX directly from bytes
                text = self.extract_text_from_docx_bytes(uploaded_file.getvalue())
            elif file_extension == '.txt':
                # Process TXT directly from bytes with proper encoding handling
                try:
                    text = uploaded_file.getvalue().decode('utf-8')
                except UnicodeDecodeError:
                    # Try other common encodings
                    for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                        try:
                            text = uploaded_file.getvalue().decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        # If all else fails, use utf-8 with error handling
                        text = uploaded_file.getvalue().decode('utf-8', errors='replace')
            else:
                st.error(f"Unsupported file format: {file_extension}")
                return None
            
            return text
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            return None
    
    def split_into_chapters(self, text: str) -> Dict[str, str]:
        """Split text into chapters based on common patterns"""
        chapters = {}
        
        # Common chapter patterns
        patterns = [
            'Chapter ',
            'CHAPTER ',
            'chapter ',
            'Ch. ',
            'CH. '
        ]
        
        # Split by chapter markers
        lines = text.split('\n')
        current_chapter = "Introduction"
        current_content = []
        chapter_num = 1
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line starts with chapter pattern
            is_chapter_header = any(line.startswith(pattern) for pattern in patterns)
            
            if is_chapter_header:
                # Save previous chapter
                if current_content:
                    chapters[current_chapter] = '\n'.join(current_content)
                
                # Start new chapter
                current_chapter = line if line else f"Chapter {chapter_num}"
                current_content = []
                chapter_num += 1
            else:
                current_content.append(line)
        
        # Save last chapter
        if current_content:
            chapters[current_chapter] = '\n'.join(current_content)
        
        return chapters
    
    def get_chapter_summary(self, chapter_text: str) -> str:
        """Generate a brief summary of the chapter"""
        # Simple summary - first 200 words
        words = chapter_text.split()
        if len(words) > 200:
            return ' '.join(words[:200]) + "..."
        return chapter_text
